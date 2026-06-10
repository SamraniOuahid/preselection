"""Partie 1 : Pages préliminaires + Chapitres I et II."""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from report_utils import *

def add_cover(doc):
    for _ in range(3): doc.add_paragraph()
    P(doc,"Université Sultan Moulay Slimane",bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,size=14,color=C1)
    P(doc,"École Nationale des Sciences Appliquées de Béni Mellal",bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,size=14,color=C1)
    doc.add_paragraph()
    P(doc,"RAPPORT DE STAGE DE FIN D'ÉTUDES",bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,size=20,color=C1)
    P(doc,FILIERE_ETUD,align=WD_ALIGN_PARAGRAPH.CENTER,size=14,color=CG)
    doc.add_paragraph(); doc.add_paragraph()
    P(doc,TITLE,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,size=16,color=C2)
    for _ in range(3): doc.add_paragraph()
    for lb,vl in [("Réalisé par",STUDENT),("Encadrant entreprise",ENCADRANT_ENT),("Encadrant académique",ENCADRANT_ACAD),("Lieu de stage","ENSA Béni Mellal — USMS"),("Période",PERIOD)]:
        pa=doc.add_paragraph(); pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r1=pa.add_run(f"{lb} : "); r1.bold=True; r1.font.size=Pt(12)
        r2=pa.add_run(vl); r2.font.size=Pt(12)
    for _ in range(3): doc.add_paragraph()
    P(doc,f"Année universitaire : {ANNEE_UNIV}",bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,size=13,color=C1)
    PB(doc)

def add_dedicace(doc):
    for _ in range(6): doc.add_paragraph()
    P(doc,"Dédicace",bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,size=18,color=C1)
    doc.add_paragraph()
    P(doc,"""À mes chers parents,
qui n'ont cessé de me soutenir et de m'encourager tout au long de mon parcours académique. Vos sacrifices, votre patience et votre amour inconditionnel sont la source de ma motivation.

À mes frères et sœurs, pour leur soutien moral et leurs encouragements constants.

À mes enseignants, qui m'ont transmis le savoir et la passion du métier d'ingénieur.

À mes amis et collègues, pour les moments de partage et d'entraide.

À tous ceux qui ont contribué, de près ou de loin, à la réalisation de ce travail.""",italic=True,align=WD_ALIGN_PARAGRAPH.CENTER,color=CG)
    PB(doc)

def add_remerciements(doc):
    doc.add_heading("Remerciements",level=1)
    P(doc,"Avant d'entamer la présentation de ce travail, je tiens à exprimer ma profonde gratitude envers toutes les personnes qui ont contribué à la réussite de ce stage de fin d'études.")
    P(doc,f"Je remercie tout d'abord mon encadrant de stage au sein de l'ENSA Béni Mellal, {ENCADRANT_ENT}, pour son encadrement rigoureux, ses orientations pertinentes et sa disponibilité tout au long de la période de stage.")
    P(doc,f"J'adresse également mes sincères remerciements à mon encadrant académique, {ENCADRANT_ACAD}, pour son suivi pédagogique et ses recommandations précieuses.")
    P(doc,"Ma reconnaissance va aussi à l'ensemble du personnel du service de la Scolarité de l'ENSA BM, qui a facilité mon intégration et m'a fourni les informations nécessaires à la compréhension du processus de présélection.")
    P(doc,"Enfin, je remercie l'USMS et l'EST de Fès pour la qualité de la formation dispensée en Bachelor Génie Logiciel.")
    PB(doc)

def add_resume(doc):
    doc.add_heading("Résumé",level=1)
    P(doc,"Le présent rapport rend compte du stage de fin d'études effectué au sein de l'ENSA Béni Mellal, rattachée à l'Université Sultan Moulay Slimane. Ce stage s'inscrit dans le cadre de l'obtention du diplôme de Bachelor en Génie Logiciel de l'EST de Fès.")
    P(doc,"Le projet consiste en la conception et le développement d'une plateforme web intelligente dédiée à la présélection des candidats aux filières d'ingénieur de l'ENSA BM. Cette plateforme automatise le traitement des dossiers de candidature, depuis le dépôt en ligne jusqu'au classement final, en passant par la vérification des documents, le scoring académique et la gestion des épreuves écrites.")
    P(doc,"La solution repose sur une architecture Full Stack moderne : Django REST Framework pour le backend et React 18 avec Tailwind CSS pour le frontend. Elle intègre un moteur de règles de rejet automatique, un algorithme de scoring pondéré conforme au système LMD marocain, un système de notifications en temps réel via WebSocket, et un module d'import des notes d'épreuve écrite au format Excel.")
    P(doc,"Mots-clés : Présélection, ENSA, Django, React, Scoring automatique, WebSocket, LMD, Full Stack.",bold=True)
    PB(doc)
    doc.add_heading("Abstract",level=1)
    P(doc,"This report presents the internship carried out at the National School of Applied Sciences of Béni Mellal (ENSA BM), part of Sultan Moulay Slimane University. The project involves the design and development of an intelligent web platform for candidate preselection to ENSA BM engineering programs.")
    P(doc,"The solution is built on a modern Full Stack architecture combining Django REST Framework for the backend and React 18 with Tailwind CSS for the frontend. It features an automatic rejection rule engine, a weighted scoring algorithm compliant with the Moroccan LMD system, real-time WebSocket notifications, and an Excel-based exam score import module.")
    P(doc,"Keywords: Preselection, ENSA, Django, React, Automatic scoring, WebSocket, LMD system, Full Stack.",bold=True)
    PB(doc)

def add_abbreviations(doc):
    doc.add_heading("Liste des abréviations",level=1)
    T(doc,["Abréviation","Signification"],[
        ("ENSA","École Nationale des Sciences Appliquées"),("USMS","Université Sultan Moulay Slimane"),
        ("EST","École Supérieure de Technologie"),("LMD","Licence — Master — Doctorat"),
        ("API","Application Programming Interface"),("REST","Representational State Transfer"),
        ("DRF","Django REST Framework"),("JWT","JSON Web Token"),
        ("UML","Unified Modeling Language"),("MCD","Modèle Conceptuel de Données"),
        ("MLD","Modèle Logique de Données"),("RBAC","Role-Based Access Control"),
        ("OCR","Optical Character Recognition"),("SMTP","Simple Mail Transfer Protocol"),
        ("CSRF","Cross-Site Request Forgery"),("XSS","Cross-Site Scripting"),
        ("TDI","Télécommunications et Développement Informatique"),
        ("IACS","Informatique Appliquée et Cybersécurité"),
        ("IAA","Ingénierie Agro-Alimentaire"),
        ("G2ER","Génie de l'Énergie et Énergies Renouvelables"),
        ("CIN","Carte d'Identité Nationale"),("CNE","Code National de l'Étudiant"),
    ])
    PB(doc)

def add_intro_generale(doc):
    doc.add_heading("Introduction générale",level=1)
    P(doc,"Dans un contexte de transformation numérique des établissements d'enseignement supérieur au Maroc, la gestion des processus d'admission et de présélection constitue un enjeu majeur pour les écoles d'ingénieurs. L'ENSA Béni Mellal accueille chaque année un nombre croissant de candidatures pour ses filières d'ingénieur, tant au niveau Bac+2 qu'au niveau Bac+3.")
    P(doc,"Le processus actuel de présélection, principalement manuel et basé sur des fichiers Excel, présente plusieurs limites : risque d'erreurs humaines, lenteur du traitement, absence de traçabilité et difficulté de gestion des volumes importants de dossiers. Le service de la Scolarité a exprimé le besoin d'une solution informatique capable d'automatiser et de fiabiliser ce processus.")
    P(doc,"C'est dans ce contexte que s'inscrit notre projet : la conception et le développement d'une plateforme web intelligente de présélection des candidats. Ce rapport est structuré en sept chapitres :")
    B(doc,[
        "Chapitre I : Présentation de l'organisme d'accueil",
        "Chapitre II : Contexte général du projet",
        "Chapitre III : Analyse et spécification des besoins",
        "Chapitre IV : Conception de la solution",
        "Chapitre V : Technologies et environnement de développement",
        "Chapitre VI : Réalisation de l'application",
        "Chapitre VII : Tests et validation",
    ])
    PB(doc)

# ═══════════════ CHAPITRE I ═══════════════

def add_ch1(doc):
    doc.add_heading("Chapitre I : Présentation de l'organisme d'accueil",level=1)

    doc.add_heading("I.1 Introduction",level=2)
    P(doc,"Ce chapitre présente l'organisme d'accueil — l'ENSA Béni Mellal —, son rattachement institutionnel, ses missions, son offre de formation, et l'environnement dans lequel s'est déroulé notre stage.")

    doc.add_heading("I.2 Présentation de l'ENSA Béni Mellal",level=2)
    doc.add_heading("I.2.1 Fiche signalétique",level=3)
    T(doc,["Rubrique","Information"],[
        ("Dénomination","École Nationale des Sciences Appliquées de Béni Mellal"),
        ("Sigle","ENSA BM"),("Université","Université Sultan Moulay Slimane (USMS)"),
        ("Statut","Établissement public d'enseignement supérieur"),
        ("Adresse","Avenue de l'Université, BP 592, Béni Mellal"),
        ("Téléphone","+212 5 23 48 11 23"),("Site web","www.ensabm.usms.ac.ma"),
        ("Création","2003 — Décret n° 2-03-200"),
        ("Effectif étudiants","~2 500 (2024-2025)"),("Enseignants","~120 enseignants-chercheurs"),
    ])

    doc.add_heading("I.2.2 Historique et création",level=3)
    P(doc,"L'ENSA BM a été créée en 2003 par décret ministériel, dans le cadre de l'extension du réseau des ENSA à travers le Maroc. Elle fait partie des douze ENSA marocaines et forme des ingénieurs d'État dans les domaines des sciences appliquées. Depuis sa création, l'école n'a cessé de se développer en diversifiant son offre de formation et en renforçant ses infrastructures.")

    doc.add_heading("I.2.3 Localisation géographique",level=3)
    P(doc,"Située à Béni Mellal, chef-lieu de la région Béni Mellal-Khénifra, l'ENSA BM est implantée sur le campus universitaire de l'USMS, favorisant les synergies avec les autres établissements de l'université.")
    FIG(doc,"Figure I.1 : Localisation géographique de l'ENSA Béni Mellal")

    doc.add_heading("I.3 Rattachement institutionnel",level=2)
    doc.add_heading("I.3.1 Université Sultan Moulay Slimane (USMS)",level=3)
    P(doc,"L'USMS, fondée en 1978, est un établissement public d'enseignement supérieur comptant plus de 80 000 étudiants répartis dans une dizaine d'établissements couvrant les sciences, lettres, droit, économie et technologies.")

    doc.add_heading("I.3.2 Position de l'ENSA BM au sein de l'USMS",level=3)
    P(doc,"L'ENSA BM occupe une position stratégique comme seul établissement de l'USMS dédié à la formation d'ingénieurs d'État, contribuant activement à la recherche scientifique et au rayonnement de l'université.")

    doc.add_heading("I.4 Missions et objectifs de l'ENSA BM",level=2)
    B(doc,[
        "Former des ingénieurs d'État polyvalents dotés de compétences scientifiques, techniques et managériales",
        "Développer la recherche scientifique et l'innovation technologique",
        "Assurer le transfert de technologie vers le tissu économique régional et national",
        "Contribuer au développement socio-économique de la région Béni Mellal-Khénifra",
        "Promouvoir la coopération nationale et internationale",
    ])

    doc.add_heading("I.5 Offre de formation",level=2)
    doc.add_heading("I.5.1 Cycle Ingénieur Bac+2",level=3)
    P(doc,"Le cycle ingénieur Bac+2 est ouvert aux titulaires d'un DUT, BTS, DEUG, DEUST ou diplôme équivalent. La sélection se fait sur étude de dossier suivie d'une épreuve écrite et/ou orale.")

    doc.add_heading("I.5.2 Cycle Ingénieur Bac+3",level=3)
    P(doc,"Le cycle Bac+3 est destiné aux titulaires d'une Licence (système LMD). Il permet l'intégration directe en première année du cycle ingénieur.")

    doc.add_heading("I.5.3 Filières proposées",level=3)
    T(doc,["Code","Filière","Niveau","Places"],[
        ("TDI","Télécommunications et Développement Informatique","Bac+2 / Bac+3","30"),
        ("IACS","Informatique Appliquée et Cybersécurité","Bac+2 / Bac+3","30"),
        ("IAA","Ingénierie Agro-Alimentaire","Bac+2 / Bac+3","30"),
        ("G2ER","Génie de l'Énergie et Énergies Renouvelables","Bac+2 / Bac+3","30"),
    ])
    FIG(doc,"Tableau I.1 : Filières d'ingénieur proposées par l'ENSA BM")

    doc.add_heading("I.6 Organisation interne",level=2)
    doc.add_heading("I.6.1 Organigramme général",level=3)
    P(doc,"L'ENSA BM est dirigée par un Directeur assisté d'un Directeur-Adjoint. L'organisation comprend plusieurs départements pédagogiques, un service de la Scolarité, un service des Affaires Étudiantes, un service Informatique et des laboratoires de recherche.")
    FIG(doc,"Figure I.2 : Organigramme de l'ENSA Béni Mellal")

    doc.add_heading("I.6.2 Service de la Scolarité",level=3)
    P(doc,"Le service de la Scolarité gère l'ensemble des processus administratifs liés aux étudiants : inscriptions, examens, stages, diplomation et, dans notre cas, la présélection des candidats au cycle ingénieur. C'est le principal bénéficiaire de notre plateforme.")

    doc.add_heading("I.7 Environnement du stage",level=2)
    doc.add_heading("I.7.1 Service d'accueil",level=3)
    P(doc,"Le stage s'est déroulé au sein du service de la Scolarité de l'ENSA BM, en collaboration étroite avec l'équipe responsable du processus de présélection des candidatures aux filières d'ingénieur.")

    doc.add_heading("I.7.2 Encadrant de stage",level=3)
    P(doc,f"Notre encadrant au sein de l'ENSA BM, {ENCADRANT_ENT}, nous a accompagné tout au long du stage en définissant les objectifs, en validant les choix techniques et en assurant le suivi de l'avancement du projet.")

    doc.add_heading("I.7.3 Conditions de travail et outils mis à disposition",level=3)
    T(doc,["Ressource","Description"],[
        ("Poste de travail","Dell Latitude 7480 — Ubuntu 24.04 LTS"),
        ("Accès réseau","Connexion Internet haut débit du campus"),
        ("Logiciels","VS Code, Git, PostgreSQL, Python 3.10, Node.js 20"),
        ("Documentation","Accès aux données anonymisées du processus de présélection"),
    ])

    doc.add_heading("I.8 Missions confiées",level=2)
    doc.add_heading("I.8.1 Contexte de la mission",level=3)
    P(doc,"Le service de la Scolarité de l'ENSA BM traite chaque année plusieurs centaines de dossiers de candidature pour les filières d'ingénieur. Le processus actuel, essentiellement manuel, nécessite une modernisation pour gagner en efficacité et en fiabilité.")

    doc.add_heading("I.8.2 Tâches réalisées",level=3)
    B(doc,[
        "Analyse du processus existant de présélection et identification des besoins",
        "Conception de l'architecture technique de la plateforme (UML, modèle de données)",
        "Développement du backend Django REST Framework avec moteur de règles et scoring",
        "Développement du frontend React avec formulaires multi-étapes et dashboard",
        "Intégration du système de notifications en temps réel via WebSocket",
        "Mise en place de l'import des notes d'épreuve écrite au format Excel",
        "Tests unitaires, d'intégration et fonctionnels",
    ])

    doc.add_heading("I.9 Conclusion",level=2)
    P(doc,"Ce chapitre nous a permis de situer le cadre institutionnel et organisationnel de notre stage. L'ENSA BM, établissement de référence au sein de l'USMS, nous a offert un environnement propice à la réalisation de notre projet. Le chapitre suivant exposera le contexte général du projet et la problématique identifiée.")
    PB(doc)

# ═══════════════ CHAPITRE II ═══════════════

def add_ch2(doc):
    doc.add_heading("Chapitre II : Contexte général du projet",level=1)

    doc.add_heading("II.1 Introduction",level=2)
    P(doc,"Ce chapitre expose le contexte général du projet, la problématique identifiée, les objectifs fixés, le périmètre fonctionnel et le planning prévisionnel de réalisation.")

    doc.add_heading("II.2 Présentation du projet",level=2)
    doc.add_heading("II.2.1 Intitulé officiel du projet",level=3)
    P(doc,f"Le projet s'intitule : « {TITLE} ». Il s'agit d'une application web Full Stack destinée à moderniser le processus de présélection des candidats aux filières d'ingénieur de l'ENSA Béni Mellal.")

    doc.add_heading("II.2.2 Commanditaire et bénéficiaires",level=3)
    T(doc,["Acteur","Rôle"],[
        ("Service Scolarité ENSA BM","Commanditaire et bénéficiaire principal"),
        ("Responsables de filières","Utilisateurs du tableau de bord"),
        ("Candidats","Utilisateurs finaux (dépôt et suivi)"),
        ("Administration ENSA BM","Supervision globale"),
    ])

    doc.add_heading("II.3 Contexte métier",level=2)
    doc.add_heading("II.3.1 Le processus de présélection actuel",level=3)
    P(doc,"Le processus actuel de présélection à l'ENSA BM se déroule en plusieurs étapes : (1) réception des dossiers physiques ou par email, (2) vérification manuelle de la conformité des pièces, (3) saisie des notes dans un fichier Excel, (4) tri et classement manuel, (5) convocation à l'épreuve écrite, (6) saisie des notes d'épreuve et classement final. Ce processus mobilise plusieurs agents du service Scolarité pendant plusieurs semaines.")

    doc.add_heading("II.3.2 Les acteurs impliqués",level=3)
    B(doc,[
        "Chef du service Scolarité : supervision du processus",
        "Agents de la Scolarité : réception, tri et saisie des dossiers",
        "Responsables de filières : validation des critères d'admissibilité",
        "Candidats : dépôt des dossiers et suivi des résultats",
    ])

    doc.add_heading("II.3.3 Le système LMD marocain et ses contraintes",level=3)
    P(doc,"Le système LMD marocain structure les études supérieures en trois cycles : Licence (6 semestres), Master (4 semestres) et Doctorat. Chaque semestre est évalué sur 20 points. Les mentions sont attribuées selon les seuils suivants :")
    T(doc,["Moyenne","Mention"],[
        ("≥ 16.00","Très Bien"),("≥ 14.00","Bien"),("≥ 12.00","Assez Bien"),("≥ 10.00","Passable"),
    ])
    P(doc,"Notre plateforme doit impérativement respecter ces seuils pour le calcul automatique des mentions et du scoring.")

    doc.add_heading("II.4 Problématique",level=2)
    doc.add_heading("II.4.1 Problèmes du processus manuel",level=3)
    B(doc,[
        "Traitement lent : plusieurs semaines pour traiter quelques centaines de dossiers",
        "Erreurs de saisie : risque élevé lors de la transcription manuelle des notes",
        "Absence de traçabilité : aucun historique des décisions prises",
        "Manque de transparence : les candidats n'ont pas de visibilité sur l'état de leur dossier",
        "Doublons non détectés : difficile d'identifier les candidatures multiples",
        "Classement approximatif : calculs manuels sujets à erreur",
    ])

    doc.add_heading("II.4.2 Risques et limites identifiés",level=3)
    T(doc,["Risque","Impact","Probabilité"],[
        ("Erreur de classement","Élevé","Moyenne"),
        ("Perte de dossiers","Critique","Faible"),
        ("Retard dans la publication","Élevé","Élevée"),
        ("Réclamations candidats","Moyen","Élevée"),
        ("Non-détection de fraudes","Élevé","Moyenne"),
    ])

    doc.add_heading("II.4.3 Besoins exprimés par le service scolarité",level=3)
    P(doc,"Les entretiens avec le service Scolarité ont permis d'identifier les besoins suivants : automatisation du tri des dossiers, détection automatique des doublons, calcul automatique du score et du classement, notification des candidats par email, traçabilité complète des décisions, et génération de rapports et exports Excel.")

    doc.add_heading("II.5 Objectifs du projet",level=2)
    doc.add_heading("II.5.1 Objectif principal",level=3)
    P(doc,"Concevoir et développer une plateforme web complète permettant d'automatiser et de fiabiliser le processus de présélection des candidats aux filières d'ingénieur de l'ENSA Béni Mellal.")

    doc.add_heading("II.5.2 Objectifs spécifiques",level=3)
    B(doc,[
        "Automatisation du rejet des dossiers non conformes via un moteur de règles configurable",
        "Réduction significative de la charge administrative du service Scolarité",
        "Classement automatique des candidats selon un algorithme de scoring pondéré",
        "Traçabilité et audit complet du processus (historique des actions)",
    ])

    doc.add_heading("II.6 Périmètre du projet",level=2)
    doc.add_heading("II.6.1 Fonctionnalités incluses",level=3)
    B(doc,[
        "Inscription et authentification sécurisée (JWT)",
        "Dépôt de dossier en ligne avec formulaire multi-étapes",
        "Upload et traitement des documents (diplôme, relevé, CIN, photo)",
        "Moteur de règles de rejet automatique configurable",
        "Algorithme de scoring avec pondération par filière",
        "Tableau de bord responsable avec KPIs et graphiques",
        "Import des notes d'épreuve écrite au format Excel",
        "Notifications en temps réel via WebSocket et email SMTP",
        "Export des résultats en Excel formaté",
    ])

    doc.add_heading("II.6.2 Fonctionnalités hors scope",level=3)
    B(doc,[
        "Intégration directe avec la plateforme nationale Massar",
        "Application mobile native",
        "Système de paiement en ligne",
        "Génération automatique des emplois du temps d'examen",
    ])

    doc.add_heading("II.7 Étude comparative des solutions existantes",level=2)
    doc.add_heading("II.7.1 Solutions similaires au Maroc",level=3)
    P(doc,"Plusieurs établissements marocains utilisent des solutions de gestion des candidatures, parmi lesquelles : la plateforme nationale Massar (gestion des bacheliers), les systèmes internes des ENSA (principalement basés sur Excel), et quelques solutions propriétaires développées localement.")

    doc.add_heading("II.7.2 Analyse comparative",level=3)
    T(doc,["Critère","Massar","Excel Manuel","Notre Solution"],[
        ("Automatisation","Partielle","Aucune","Complète"),
        ("Personnalisation","Limitée","Manuelle","Totale par filière"),
        ("Scoring automatique","Non","Non","Oui (configurable)"),
        ("Notifications temps réel","Non","Non","Oui (WebSocket)"),
        ("Traçabilité","Limitée","Aucune","Complète"),
        ("Import notes Excel","Non","Manuel","Automatisé"),
        ("Coût","Gratuit","Gratuit","Développement interne"),
    ])
    FIG(doc,"Tableau II.1 : Analyse comparative des solutions existantes")

    doc.add_heading("II.7.3 Positionnement de notre solution",level=3)
    P(doc,"Notre solution se positionne comme une plateforme sur mesure, conçue spécifiquement pour le contexte de l'ENSA BM. Elle offre un niveau d'automatisation et de personnalisation supérieur aux solutions existantes, tout en restant modulaire et extensible.")

    doc.add_heading("II.8 Planning prévisionnel du projet",level=2)
    T(doc,["Semaine","Phase","Livrables"],[
        ("S1","Fondations Backend","Modèles Django, API REST CRUD, Authentification JWT"),
        ("S2","Intelligence et logique métier","Moteur de règles, Scoring, Extraction documents"),
        ("S3","Frontend React","Pages candidat, Dashboard responsable, Formulaires"),
        ("S4","Tests et déploiement","Tests unitaires/intégration, Corrections, Documentation"),
    ])
    FIG(doc,"Tableau II.2 : Planning prévisionnel du projet (diagramme de Gantt)")

    doc.add_heading("II.9 Conclusion",level=2)
    P(doc,"Ce chapitre a permis de cerner le contexte du projet, la problématique du processus manuel de présélection et les objectifs de la solution proposée. Le chapitre suivant détaillera l'analyse et la spécification des besoins fonctionnels et non fonctionnels.")
    PB(doc)
