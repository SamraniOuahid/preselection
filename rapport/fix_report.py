#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import RGBColor

def corriger_styles_et_couleurs():
    # Prends le fichier généré par le script principal
    input_file = "Rapport_EST_Fes_SAMRANI_Ouahid.docx"
    output_file = "Rapport_EST_Fes_FINAL_NOIR.docx"
    
    if not os.path.exists(input_file):
        print(f"❌ Erreur : Le fichier {input_file} est introuvable.")
        return

    print("==================================================")
    print("   CORRECTION STRICTE DU RAPPORT (EST FÈS)        ")
    print("==================================================")
    
    doc = Document(input_file)
    COLOR_BLACK = RGBColor(0, 0, 0)
    FONT_NAME = "Latin Modern Roman"

    # 1. Correction de tous les paragraphes et titres
    print("-> Forçage de la police et du noir sur le texte...")
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.color.rgb = COLOR_BLACK

    # 2. DESTRUCTION DU STYLE BLEU DES TABLEAUX
    print("-> Forçage du style 'Table Grid' (Noir et Blanc) sur les tableaux...")
    for table in doc.tables:
        # Appliquer la grille noire classique de Word
        table.style = 'Table Grid'
        
        # Nettoyer l'intérieur de chaque cellule au cas où
        for row in table.rows:
            for cell in row.cells:
                # Retirer tout ombrage résiduel
                tcPr = cell._tc.get_or_add_tcPr()
                for child in list(tcPr):
                    if child.tag.endswith('shd'):
                        tcPr.remove(child)
                
                # Forcer le texte à l'intérieur du tableau en noir et Latin Modern Roman
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = FONT_NAME
                        run.font.color.rgb = COLOR_BLACK

    # Sauvegarde
    doc.save(output_file)
    print("==================================================")
    print("✅ Opération réussie !")
    print(f"📄 Fichier 100% propre : {output_file}")
    print("==================================================")

if __name__ == "__main__":
    corriger_styles_et_couleurs()