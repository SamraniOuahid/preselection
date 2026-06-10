"""
Configuration et fonctions utilitaires pour le rapport EST Fès.
Police : Latin Modern Roman — Typographie stricte EST Fès.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════
# VARIABLES DE CONFIGURATION — À MODIFIER PAR L'ÉTUDIANT
# ══════════════════════════════════════════════════════════════
NOM_ETUDIANT = "SAMRANI Ouahid"
THEME_STAGE = (
    "Conception et Développement d'une Plateforme Intelligente "
    "de Présélection des Candidats"
)
ENCADRANT_ACADEMIQUE = "[À REMPLIR : Nom Encadrant EST Fès]"
ENCADRANT_ENTREPRISE = "[À REMPLIR : Nom Encadrant ENSA BM]"
JURY_1 = "[À REMPLIR : Nom Membre Jury 1]"
JURY_2 = "[À REMPLIR : Nom Membre Jury 2]"
DATES_STAGE = "[À REMPLIR : Du jj/mm/aaaa au jj/mm/aaaa]"
ANNEE_UNIV = "2025-2026"
LIEU_STAGE = "ENSA Béni Mellal — USMS"
FILIERE = "Génie Logiciel"
DIPLOME = "Bachelor Universitaire de Technologie"
ETABLISSEMENT = "EST Fès — Université Sidi Mohamed Ben Abdellah"

# ══════════════════════════════════════════════════════════════
# COULEURS
# ══════════════════════════════════════════════════════════════
C_DARK = RGBColor(0x1A, 0x1A, 0x2E)
C_ACCENT = RGBColor(0x00, 0x56, 0x8A)
C_BODY = RGBColor(0x33, 0x33, 0x33)
C_GRAY = RGBColor(0x66, 0x66, 0x66)

FONT_NAME = "Latin Modern Roman"

# ══════════════════════════════════════════════════════════════
# SETUP DOCUMENT
# ══════════════════════════════════════════════════════════════

def setup_document():
    """Crée et configure le document selon les normes EST Fès."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Style Normal : 12pt Latin Modern Roman
    normal = doc.styles['Normal']
    normal.font.name = FONT_NAME
    normal.font.size = Pt(12)
    normal.font.color.rgb = C_BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5

    # Heading 1 : Titre de Chapitre / Introduction / Conclusion — 18pt
    h1 = doc.styles['Heading 1']
    h1.font.name = FONT_NAME
    h1.font.size = Pt(18)
    h1.font.color.rgb = C_DARK
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2 : Titre de Section (ex: 1.1) — 14pt
    h2 = doc.styles['Heading 2']
    h2.font.name = FONT_NAME
    h2.font.size = Pt(14)
    h2.font.color.rgb = C_ACCENT
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3 : Titre de Sous-section — 12pt
    h3 = doc.styles['Heading 3']
    h3.font.name = FONT_NAME
    h3.font.size = Pt(12)
    h3.font.color.rgb = C_ACCENT
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    return doc

# ══════════════════════════════════════════════════════════════
# FONCTIONS UTILITAIRES
# ══════════════════════════════════════════════════════════════

def P(doc, text, bold=False, italic=False, align=None, size=None, color=None):
    """Ajoute un paragraphe formaté."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT_NAME
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        para.alignment = align
    return para


def B(doc, items):
    """Ajoute une liste à puces."""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME


def T(doc, headers, rows, caption=None):
    """Ajoute un tableau formaté avec caption optionnelle."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # En-têtes
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = str(h)
        for pr in cell.paragraphs:
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in pr.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = FONT_NAME
    # Données
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for pr in cell.paragraphs:
                for r in pr.runs:
                    r.font.size = Pt(10)
                    r.font.name = FONT_NAME
    doc.add_paragraph()
    if caption:
        FIG(doc, caption)
    return tbl


def FIG(doc, caption):
    """Ajoute une légende de figure/tableau."""
    P(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
      color=C_GRAY, size=10)


def PB(doc):
    """Ajoute un saut de page."""
    doc.add_page_break()


def chapter_title(doc, text):
    """Ajoute un titre de chapitre (Heading 1) avec saut de page avant."""
    doc.add_heading(text, level=1)


def section_title(doc, text):
    """Ajoute un titre de section (Heading 2)."""
    doc.add_heading(text, level=2)


def subsection_title(doc, text):
    """Ajoute un titre de sous-section (Heading 3)."""
    doc.add_heading(text, level=3)
