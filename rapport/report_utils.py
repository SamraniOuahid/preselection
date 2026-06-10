"""Utilitaires partagés pour la génération du rapport de stage."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TITLE = "Conception et Développement d'une Plateforme Intelligente de Présélection des Candidats"
STUDENT = "SAMRANI Ouahid"
SCHOOL = "ENSA Béni Mellal"
FILIERE_ETUD = "Bachelor Génie Logiciel — EST Fès"
PERIOD = "Mai — Juin 2026"
ENCADRANT_ENT = "Pr. [Nom Encadrant ENSA BM]"
ENCADRANT_ACAD = "Pr. [Nom Encadrant EST Fès]"
ANNEE_UNIV = "2025 — 2026"

C1 = RGBColor(0x00, 0x3D, 0x6B)
C2 = RGBColor(0x1A, 0x73, 0xE8)
C3 = RGBColor(0xE8, 0x6C, 0x00)
CT = RGBColor(0x33, 0x33, 0x33)
CG = RGBColor(0x66, 0x66, 0x66)

def setup_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3); s.right_margin = Cm(2.5)
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'; st.font.size = Pt(12); st.font.color.rgb = CT
    st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.5
    for i, (sz, col) in enumerate([(18,C1),(16,C1),(14,C2),(13,C3)], 1):
        h = doc.styles[f'Heading {i}']
        h.font.name='Arial'; h.font.size=Pt(sz); h.font.color.rgb=col; h.font.bold=True
        h.paragraph_format.space_before = Pt(18 if i<=2 else 12)
        h.paragraph_format.space_after = Pt(8)
    return doc

def P(doc, text, bold=False, italic=False, align=None, size=None, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold; run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align: para.alignment = align
    return para

def T(doc, headers, rows):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = h
        for pr in c.paragraphs:
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in pr.runs: r.bold = True; r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]; c.text = str(val)
            for pr in c.paragraphs:
                for r in pr.runs: r.font.size = Pt(10)
    doc.add_paragraph()
    return tbl

def B(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def PB(doc):
    doc.add_page_break()

def FIG(doc, caption):
    P(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=CG, size=10)
