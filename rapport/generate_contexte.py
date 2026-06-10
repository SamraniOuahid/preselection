import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def setup_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3); s.right_margin = Cm(2.5)
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'; st.font.size = Pt(12)
    st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.5
    
    C1 = RGBColor(0x00, 0x3D, 0x6B)
    C2 = RGBColor(0x1A, 0x73, 0xE8)
    
    for i, (sz, col) in enumerate([(18,C1),(16,C1),(14,C2)], 1):
        h = doc.styles[f'Heading {i}']
        h.font.name='Arial'; h.font.size=Pt(sz); h.font.color.rgb=col; h.font.bold=True
        h.paragraph_format.space_before = Pt(18 if i<=2 else 12)
        h.paragraph_format.space_after = Pt(8)
    return doc

def P(doc, text, bold=False, italic=False, align=None, size=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold; run.italic = italic
    if size: run.font.size = Pt(size)
    if align: para.alignment = align
    return para

def B(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def T(doc, headers, rows):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = str(h)
        for pr in c.paragraphs:
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in pr.runs: r.bold = True; r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]; c.text = str(val)
            for pr in c.paragraphs:
                for r in pr.runs: r.font.size = Pt(10)
    doc.add_paragraph()
    return tbl

def generate():
    doc = setup_doc()
    
    # Titre
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titre.add_run("CONTEXTE TECHNIQUE DU PROJET\nPlateforme de Présélection ENSA Béni Mellal")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x00, 0x3D, 0x6B)
    doc.add_paragraph()

    # 1. Architecture Globale
    doc.add_heading("1. Architecture Globale", level=1)
    P(doc, "L'application repose sur une architecture moderne Full-Stack et modulaire :")
    B(doc, [
        "Frontend : Développé en React.js (compilé via Vite) avec Tailwind CSS pour un rendu esthétique et responsive. L'architecture s'appuie sur des composants réutilisables, un state management optimisé et React Router.",
        "Backend : Propulsé par Django et Django REST Framework (DRF), exposant une API REST sécurisée via JWT.",
        "Base de Données : Système relationnel géré par l'ORM Django (SQLite en dev, compatible PostgreSQL en production) pour la persistance des candidatures et configurations.",
        "Temps Réel : Intégration de Django Channels avec Redis pour gérer les WebSockets (notifications en temps réel)."
    ])

    # 2. Dictionnaire des Modèles Django
    doc.add_heading("2. Dictionnaire des Modèles Django", level=1)
    P(doc, "Le modèle relationnel est structuré en plusieurs applications (apps) pour séparer les responsabilités :")
    
    doc.add_heading("Application Users", level=2)
    B(doc, [
        "User : Modèle personnalisé étendant AbstractBaseUser. Gère l'authentification (email/CIN) et les rôles (ADMIN, RESPONSABLE, CANDIDAT).",
        "Candidat : Profil détaillé de l'étudiant (relié en OneToOne à User)."
    ])
    
    doc.add_heading("Application Candidatures", level=2)
    B(doc, [
        "Dossier : Modèle central liant un Candidat à une Filiere. Contient le statut, la moyenne, les scores et motifs de rejet.",
        "Document : Pièces justificatives (Diplôme, Relevé de notes, CIN) uploadées.",
        "NoteSemestre : Évaluation par semestre du candidat (remplace l'ancienne approche par NoteMatiere), intégrant moyenne, session et mention dynamique."
    ])
    
    doc.add_heading("Application Administration & Scoring", level=2)
    B(doc, [
        "Filiere : Définit les filières du cycle ingénieur.",
        "EpreuveEcrite / NoteEcrite : Gestion de la logistique du concours écrit et des résultats d'admission.",
        "ConfigScoring & RegleRejet : Paramétrage dynamique des pondérations par matière et des règles de rejet automatiques.",
        "HistoriqueAction : Traçabilité détaillée des changements de statuts et interventions manuelles."
    ])

    # 3. Coefficients des 4 Filières
    doc.add_heading("3. Coefficients des 4 Filières", level=1)
    P(doc, "Le moteur de scoring applique les pondérations suivantes sur les matières extraites, en fonction de la filière choisie :")
    headers = ["Filière", "Pondérations"]
    rows = [
        ["TDI", "INFO: 35%, MATH: 30%, ELEC_AUTO: 25%, LANGUES: 10%"],
        ["IACS", "INFO: 40%, MATH: 35%, RESEAUX_BD: 15%, LANGUES: 10%"],
        ["IAA", "CHIMIE_BIO: 60%, MATH: 20%, INFO: 10%, LANGUES: 10%"],
        ["G2ER", "MATH: 30%, CHIMIE_BIO: 30%, ELEC_AUTO: 25%, INFO: 15%"]
    ]
    T(doc, headers, rows)

    # 4. Fonctionnement de l'OCR et de la Logique Floue
    doc.add_heading("4. Fonctionnement de l'OCR et de la Logique Floue", level=1)
    P(doc, "Le pipeline de traitement automatique des relevés de notes s'appuie sur plusieurs étapes de fiabilisation :")
    B(doc, [
        "Extraction OCR : Le texte brut est d'abord extrait des documents numérisés.",
        "Nettoyage du bruit : Une fonction spécialisée filtre les erreurs fréquentes d'OCR sur les chiffres (ex: 'O' interprété comme '0', 'l' comme '1', 'S' comme '5'). Les textes sont normalisés et désaccentués.",
        "Fuzzy Matching (Logique Floue) : L'algorithme utilise la librairie `fuzzywuzzy` pour comparer les noms de matières extraits avec un dictionnaire de catégories standardisées (MATH, INFO, CHIMIE_BIO...). Un seuil de similarité strict de 80% est appliqué.",
        "Scoring et Fallback : Les notes sont regroupées par catégorie. Si une catégorie indispensable à la filière n'est pas détectée, le moteur applique la moyenne générale (fallback) et marque automatiquement le dossier comme 'SUSPECT' pour vérification manuelle."
    ])

    # 5. Gestion de l'asynchronisme et des logs
    doc.add_heading("5. Gestion de l'Asynchronisme et des Logs", level=1)
    
    doc.add_heading("Asynchronisme", level=2)
    P(doc, "Pour garantir une excellente expérience utilisateur lors des opérations coûteuses (ex: traitement OCR, calculs de score, envois d'e-mails), l'application utilise l'asynchronisme via `threading.Thread`. Les processus lourds sont exécutés en arrière-plan (threads démons). Le système veille également à clore correctement les connexions à la base de données dans les threads pour éviter toute fuite ou conflit.")
    
    doc.add_heading("Système de Logging", level=2)
    P(doc, "L'application intègre un mécanisme de traçabilité sécurisé par fichiers. La configuration Django dirige les flux vers un fichier central (`logs/ensa_presel.log`). Les logs sont structurés par domaine d'application (`django`, `candidatures`, `notifications`, `channels`) avec des niveaux de criticité adaptés (INFO, ERROR, DEBUG), permettant un diagnostic rapide en production.")

    # Sauvegarde
    out_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Contexte_Technique_Projet.docx")
    doc.save(out_path)
    print(f"Fichier généré avec succès à l'emplacement : {out_path}")

if __name__ == "__main__":
    generate()
