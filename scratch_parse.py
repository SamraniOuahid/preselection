import docx

doc_path = "/home/ouahid/myproject/stage/preseclection/rapport/Rapport_EST_Fes_FINAL_NOIR.docx"
try:
    doc = docx.Document(doc_path)
    print("Loaded document.")
    print("Number of paragraphs:", len(doc.paragraphs))
    print("Number of tables:", len(doc.tables))
    for i, p in enumerate(doc.paragraphs[:20]):
        print(f"P{i} (style={p.style.name}): {p.text}")
except Exception as e:
    print("Error:", e)
